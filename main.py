import os
import re
import json
import time
import math
import unicodedata
from typing import List, Dict, Any
from urllib.parse import urlparse
from pathlib import Path

try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False

import langid
from langchain_ollama import ChatOllama
from langchain_core.messages import HumanMessage

from config import ModelConfig, SYSTEM_PROMPT
from schemas import AnalysisOutput

# Initialize global variables
TRAFILATURA_AVAILABLE = False
PSUTIL_AVAILABLE = False
PYTESSERACT_AVAILABLE = False
PIL_AVAILABLE = False
PYPDF_AVAILABLE = False
PYTHON_DOCX_AVAILABLE = False
trafilatura = None

try:
    import trafilatura
    TRAFILATURA_AVAILABLE = True
except ImportError:
    TRAFILATURA_AVAILABLE = False

try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False

try:
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = r'D:/Project/llm_clean/tesseract/tesseract.exe'  # Ensure tesseract is in PATH

    from PIL import Image

    PIL_AVAILABLE = True
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False
    PIL_AVAILABLE = False

try:
    import PyPDF2
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# Single LLM instance
llm = None

def get_llm():
    """Get or create LLM instance"""
    global llm
    if llm is None:
        try:
            llm = ChatOllama(
                model=ModelConfig.MODEL_NAME,
                temperature=ModelConfig.TEMPERATURE,
                num_ctx=ModelConfig.CONTEXT_WINDOW,
                top_p=ModelConfig.TOP_P,
                top_k=ModelConfig.TOP_K,
                keep_alive="5m"
            )
        except Exception as e:
            print(f"Error initializing LLM: {str(e)}")
            return None
    return llm


# ======================================================
# DOCUMENT AND IMAGE EXTRACTION
# ======================================================

def extract_text_from_image(image_path: str) -> str:
    """Extract text from image using OCR"""
    if not PYTESSERACT_AVAILABLE or not PIL_AVAILABLE:
        return "Error: pytesseract and PIL not installed. Install: pip install pytesseract pillow"
    
    try:
        if not os.path.exists(image_path):
            return f"Error: Image file not found: {image_path}"
        
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image)
        
        if not text.strip():
            return f"Error: No text found in image {image_path}"
        
        return text.strip()
    except Exception as e:
        return f"Error extracting text from image: {str(e)}"


def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF file"""
    if not PYPDF_AVAILABLE:
        return "Error: PyPDF2 not installed. Install: pip install PyPDF2"
    
    try:
        if not os.path.exists(pdf_path):
            return f"Error: PDF file not found: {pdf_path}"
        
        text_content = []
        
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            if num_pages == 0:
                return "Error: PDF file is empty"
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    text_content.append(text)
        
        if not text_content:
            return f"Error: No text found in PDF {pdf_path}"
        
        return "\n\n".join(text_content).strip()
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"


def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from DOCX file"""
    if not PYTHON_DOCX_AVAILABLE:
        return "Error: python-docx not installed. Install: pip install python-docx"
    
    try:
        if not os.path.exists(docx_path):
            return f"Error: DOCX file not found: {docx_path}"
        
        doc = Document(docx_path)
        text_content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        if not text_content:
            return f"Error: No text found in DOCX {docx_path}"
        
        return "\n\n".join(text_content).strip()
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"


def extract_text_from_file(file_path: str) -> str:
    """Extract text from any supported file format"""
    if not os.path.exists(file_path):
        return f"Error: File not found: {file_path}"
    
    file_ext = Path(file_path).suffix.lower()
    
    try:
        if file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif']:
            return extract_text_from_image(file_path)
        elif file_ext == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return extract_text_from_docx(file_path)
        elif file_ext in ['.txt', '.text']:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read().strip()
                if text:
                    return text
                return "Error: TXT file is empty"
        else:
            return f"Error: Unsupported file format: {file_ext}. Supported: JPG, PNG, PDF, DOCX, TXT"
    except Exception as e:
        return f"Error extracting text: {str(e)}"


# ======================================================
# UTILITIES
# ======================================================

def clean_text(text: str) -> str:
    """Clean and normalize text"""
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"http[s]?://\S+|www\.\S+", " ", text)
    text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def is_url(text: str) -> bool:
    """Check if input is a URL"""
    try:
        result = urlparse(text.strip())
        return all([result.scheme, result.netloc])
    except Exception:
        return False


def is_file_path(text: str) -> bool:
    """Check if input is a file path"""
    try:
        return os.path.exists(text.strip())
    except Exception:
        return False


def fetch_url_content(url: str) -> str:
    """Fetch and extract text content from URL using trafilatura"""
    if not TRAFILATURA_AVAILABLE:
        return f"Error: trafilatura not installed"
    
    try:
        downloaded = trafilatura.fetch_url(url)
        
        if downloaded is None:
            return f"Error: Could not fetch content from {url} - check URL or network connection"
        
        extracted_text = trafilatura.extract(downloaded, include_tables=True, include_formatting=False, no_fallback=True, include_images=False, include_links=False)
        
        if extracted_text is None or len(extracted_text.strip()) == 0:
            return f"Error: Could not extract text from {url} - page may be empty or blocked"
        
        return extracted_text.strip()
    except Exception as e:
        return f"Error: {str(e)}"


def extract_json(text: str) -> dict:
    """Extract JSON from text, handling malformed JSON"""
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("No valid JSON found")
    
    json_str = text[start:end + 1]
    
    try:
        return json.loads(json_str)
    except json.JSONDecodeError:
        json_str = re.sub(r',(\s*[}\]])', r'\1', json_str)
        
        try:
            return json.loads(json_str)
        except json.JSONDecodeError:
            json_str = json_str.replace('\n', '\\n').replace('\r', '\\r').replace('\t', '\\t')
            
            try:
                return json.loads(json_str)
            except json.JSONDecodeError:
                brace_count = 0
                in_string = False
                escape_next = False
                valid_end = -1
                
                for i, char in enumerate(json_str[start:]):
                    if escape_next:
                        escape_next = False
                        continue
                    if char == '\\':
                        escape_next = True
                        continue
                    if char == '"' and not escape_next:
                        in_string = not in_string
                        continue
                    if not in_string:
                        if char == '{':
                            brace_count += 1
                        elif char == '}':
                            brace_count -= 1
                            if brace_count == 0:
                                valid_end = i + 1
                                break
                
                if valid_end > 0:
                    try:
                        return json.loads(json_str[start:start + valid_end])
                    except Exception:
                        pass
                
                raise ValueError("Cannot parse JSON after all cleanup attempts")


def detect_script(text: str) -> List[Dict[str, Any]]:
    """Detect script type and percentage usage"""
    scripts = {
        "Devanagari": (0x0900, 0x097F),
        "Tamil": (0x0B80, 0x0BFF),
        "Telugu": (0x0C00, 0x0C7F),
        "Kannada": (0x0C80, 0x0CFF),
        "Malayalam": (0x0D00, 0x0D7F),
        "Bengali": (0x0980, 0x09FF),
        "Gujarati": (0x0A80, 0x0AFF),
        "Punjabi": (0x0A00, 0x0A7F),
        "Odia": (0x0B00, 0x0B7F),
        "Roman/Latin": (0x0041, 0x005A),
    }
    
    script_counts = {name: 0 for name in scripts.keys()}
    
    for char in text:
        code = ord(char)
        for script_name, (start, end) in scripts.items():
            if start <= code <= end:
                script_counts[script_name] += 1
                break
    
    total = len([c for c in text if ord(c) > 64])
    
    detected = [(name, count) for name, count in script_counts.items() if count > 0]
    detected.sort(key=lambda x: x[1], reverse=True)
    
    result = []
    for idx, (name, count) in enumerate(detected[:2], 1):
        percentage = round((count / total * 100), 2) if total > 0 else 0
        result.append({
            f"script_detected_{idx}": name,
            f"confidence_script_{idx}": round(percentage / 100, 2)
        })
    
    return result


def get_system_metrics() -> tuple:
    """Get system metrics: gpu_utilisation and memory_usage"""
    try:
        if PSUTIL_AVAILABLE:
            memory = psutil.virtual_memory()
            memory_usage = f"{memory.percent}%"
        else:
            memory_usage = "N/A"
    except:
        memory_usage = "N/A"
    
    gpu_util = "N/A"
    
    return gpu_util, memory_usage


# ======================================================
# LANGUAGE DETECTION
# ======================================================

def detect_language(text: str) -> tuple:
    """Detect language and confidence"""
    try:
        lang, score = langid.classify(text)
        confidence = round(1 / (1 + math.exp(-score / 100)), 2)
    except Exception:
        lang = "en"
        confidence = 0.5
    
    return lang, confidence


# ======================================================
# MAIN ANALYSIS PIPELINE
# ======================================================

def analyze(input_data: str) -> dict:
    """Main analysis function"""
    start = time.perf_counter()

    original_input = input_data.strip()
    
    # Get text from file, URL, or direct input
    if is_file_path(original_input):
        # If input is a file path, extract text from file
        text = extract_text_from_file(original_input)
    elif is_url(original_input):
        # If input is a URL, fetch its content
        text = fetch_url_content(original_input)
    else:
        # If input is text, don't clean it - preserve original
        text = original_input

    # Get LLM instance
    llm_instance = get_llm()
    if not llm_instance:
        return {"error": "Failed to initialize LLM"}

    # Call main analysis with system prompt
    prompt = SYSTEM_PROMPT.replace("{INPUT_TEXT}", text[:2500])
    
    parsed = {}
    try:
        response = llm_instance.invoke([HumanMessage(content=prompt)])
        raw = response.content.strip()
        
        try:
            parsed = json.loads(raw)
        except json.JSONDecodeError:
            try:
                parsed = extract_json(raw)
            except Exception as e:
                print(f"JSON parsing error: {str(e)}")
                parsed = {}
    except Exception as e:
        print(f"LLM call failed: {str(e)}")
        parsed = {}

    if not isinstance(parsed, dict):
        parsed = {}

    # Detect language
    detected_lang, lang_confidence = detect_language(text)

    # Detect scripts
    scripts = detect_script(text)
    
    # Build language_scripts with detected language and scripts
    if not parsed.get("language_scripts"):
        language_scripts = []
        for idx, script_dict in enumerate(scripts[:2], 1):
            lang_script = {
                f"lang_detected_{idx}": detected_lang,
                f"confidence_lang": lang_confidence,
            }
            lang_script.update(script_dict)
            language_scripts.append(lang_script)
        parsed["language_scripts"] = language_scripts if language_scripts else [{}]

    # Set required fields from original input
    parsed["original_input"] = original_input

    # Set defaults
    parsed.setdefault("reasoning_summary", "")
    parsed.setdefault("translation", "")
    parsed.setdefault("transliteration", "")
    parsed.setdefault("summary", "")
    parsed.setdefault("ner", {
        "PERSON": [],
        "LOCATION": [],
        "ORGANIZATION": [],
        "PRODUCT": []
    })
    parsed.setdefault("events", [])
    parsed.setdefault("country_id", [])
    parsed.setdefault("domain", [])

    # Get system metrics
    gpu_util, memory_usage = get_system_metrics()
    parsed["gpu_utilisation"] = gpu_util
    parsed["memory_usage"] = memory_usage

    # Execution time in seconds
    parsed["execution_time_sec"] = round(
        (time.perf_counter() - start), 2
    )

    # Try to validate
    try:
        validated = AnalysisOutput(**parsed)
        return validated.model_dump()
    except Exception as e:
        print(f"Validation warning: {str(e)}")
        return parsed